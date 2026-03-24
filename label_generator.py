import streamlit as st
import pandas as pd
import datetime
import re
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pystrich.code128 import Code128Encoder
from PIL import Image, ImageDraw, ImageFont
from io import BytesIO

BARCODE_WIDTH     = 3.6
BARCODE_HEIGHT    = 1.9
BARCODE_TEXT_SIZE = 18
MAX_ROWS          = 10

XIMMIO_EXPORT_COLUMNS = {'Stad', 'Straat', 'Huisnummer', 'Postcode', 'SubTaskDesc'}

# GitHub raw URL voor het Avalex logo
AVALEX_LOGO_URL = "https://raw.githubusercontent.com/your-org/your-repo/main/Avalexlogo.png"


# -------------------------------------------------------
# Hulpfuncties
# -------------------------------------------------------

def strip_spaces(value):
    """Verwijder alle spaties uit een veld."""
    return str(value).replace(" ", "").strip()


def generate_word_from_dataframe(df):
    """Genereer Word-document vanuit een DataFrame met interne kolomnamen:
    containertype, straat, huisnummer, toevoeging, postcode, woonplaats
    """
    output_doc = Document()
    section = output_doc.sections[-1]
    section.page_width    = Cm(10.0)
    section.page_height   = Cm(8.0)
    section.left_margin   = Cm(0.8)
    section.right_margin  = Cm(0.5)
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(0.5)

    try:
        font = ImageFont.truetype("arial.ttf", BARCODE_TEXT_SIZE)
    except IOError:
        font = ImageFont.load_default()

    dpi              = 96
    pixels_per_cm    = dpi / 2.54
    horizontal_shift = int(0.75 * pixels_per_cm)
    text_area_height = 30

    for label_idx, (idx, row) in enumerate(df.iterrows()):
        containertype = str(row.get('containertype', ''))
        straat        = str(row.get('straat', ''))
        huisnummer    = str(row.get('huisnummer', ''))
        toevoeging    = str(row.get('toevoeging', ''))
        postcode      = str(row.get('postcode', ''))
        woonplaats    = str(row.get('woonplaats', ''))
        barcode_value = f"{postcode}{huisnummer}{toevoeging}"

        encoder       = Code128Encoder(barcode_value)
        barcode_img   = encoder.get_imagedata()
        barcode_image = Image.open(BytesIO(barcode_img))

        bbox = barcode_image.getbbox()
        if bbox:
            left, _, right, _ = bbox
            barcode_image = barcode_image.crop((left, 0, right, barcode_image.height))

        draw = ImageDraw.Draw(barcode_image)
        width, height = barcode_image.size
        draw.rectangle([0, height - text_area_height, width, height], fill="white")

        text      = ""
        bbox_text = draw.textbbox((0, 0), text, font=font)
        text_y    = height - text_area_height + ((text_area_height - (bbox_text[3] - bbox_text[1])) / 2)
        draw.text((horizontal_shift, text_y), text, fill="black", font=font)

        barcode_buf = BytesIO()
        barcode_image.save(barcode_buf, format="PNG")
        barcode_buf.seek(0)

        # Voeg page break toe vóór elk label (behalve het eerste)
        if label_idx > 0:
            output_doc.add_page_break()

        p_title = output_doc.add_paragraph(containertype)
        for run in p_title.runs:
            run.font.name = 'Arial'
            run.bold = True
        p_title.style.font.size = Pt(12)
        p_title.paragraph_format.space_before = Pt(0)
        p_title.paragraph_format.space_after = Pt(4)

        p_img   = output_doc.add_paragraph()
        run_img = p_img.add_run()
        run_img.add_picture(barcode_buf, width=Cm(BARCODE_WIDTH), height=Cm(BARCODE_HEIGHT))
        p_img.paragraph_format.space_before = Pt(0)
        p_img.paragraph_format.space_after = Pt(4)

        p_info1 = output_doc.add_paragraph(f"{straat} {huisnummer} {toevoeging}".strip())
        for run in p_info1.runs:
            run.font.name = 'Arial'
            run.bold = True
        p_info1.style.font.size = Pt(12)
        p_info1.paragraph_format.space_before = Pt(0)
        p_info1.paragraph_format.space_after = Pt(2)

        p_info2 = output_doc.add_paragraph(f"{postcode} {woonplaats}")
        for run in p_info2.runs:
            run.font.name = 'Arial'
            run.bold = True
        p_info2.style.font.size = Pt(12)
        p_info2.paragraph_format.space_before = Pt(0)
        p_info2.paragraph_format.space_after = Pt(0)

    docx_buffer = BytesIO()
    output_doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer


def parse_subtaskdesc(value):
    """Parseer SubTaskDesc en geef (category, containertype) terug.
    Categorieën: CHANGE, NEW, EXTRA, REMOVE (of None als niet herkend).
    """
    if not value or str(value).strip().lower() in ('', 'nan', 'none'):
        return None, None

    s = str(value).strip()
    match = re.search(r'-\s*(CHANGE|NEW|EXTRA|REMOVE)\b', s, re.IGNORECASE)
    if not match:
        return None, None

    cat = match.group(1).upper()

    if cat == 'REMOVE':
        return 'REMOVE', None

    if cat == 'CHANGE':
        # Zoek laatste (...) blok en pak containertype na laatste >
        inner_match = re.search(r'\(([^)]+)\)\s*$', s)
        if inner_match:
            inner = inner_match.group(1)
            after_arrow = re.search(r'>\s*(\S+)\s*$', inner)
            if after_arrow:
                return cat, after_arrow.group(1)
        return cat, None

    # NEW of EXTRA: containertype is de inhoud van het laatste (...) blok
    inner_match = re.search(r'\(([^)]+)\)\s*$', s)
    if inner_match:
        return cat, inner_match.group(1).strip()

    return cat, None


def is_ximmio_export(df):
    """Detecteer of het bestand een Ximmio bakwagen export is."""
    return XIMMIO_EXPORT_COLUMNS.issubset(set(df.columns))


def dataframe_from_ximmio_export(df, skip_indices=None):
    """Map Ximmio bakwagen export naar intern DataFrame formaat.
    skip_indices: set van DataFrame-indices die overgeslagen moeten worden (validatiefouten).
    """
    if skip_indices is None:
        skip_indices = set()
    rows = []
    for idx, row in df.iterrows():
        subtask = str(row.get('SubTaskDesc', ''))
        cat, container = parse_subtaskdesc(subtask)

        if cat == 'REMOVE':
            continue  # Overslaan

        if idx in skip_indices:
            continue  # Overslaan vanwege validatiefouten

        hl_raw     = row.get('Huisletter', '')
        tv_raw     = row.get('Huisnummer toevoeging', '')
        huisletter = '' if pd.isna(hl_raw) else str(hl_raw).strip()
        toevoeging = '' if pd.isna(tv_raw) else str(tv_raw).strip()

        rows.append({
            'containertype': strip_spaces(container or ''),
            'straat':        str(row.get('Straat', '')).strip(),
            'huisnummer':    str(row.get('Huisnummer', '')).strip(),
            'toevoeging':    (huisletter + toevoeging).strip(),
            'postcode':      strip_spaces(str(row.get('Postcode', ''))),
            'woonplaats':    str(row.get('Stad', '')).strip(),
            '_cat':          cat or '',
            '_zipcode_raw':  str(row.get('Postcode', '')).strip(),
            '_huisnummer_raw': str(row.get('Huisnummer', '')).strip(),
            '_huisletter_raw': huisletter,
            '_toevoeging_raw': toevoeging,
        })

    result_df = pd.DataFrame(rows) if rows else pd.DataFrame()
    return result_df


def dataframe_from_file(file):
    """Lees CSV/XLSX/CSV en detecteer automatisch het formaat.
    Geeft een tuple terug: (df, category_counts)
    """
    df_raw = pd.read_excel(file) if file.name.endswith(".xlsx") else pd.read_csv(file)

    if is_ximmio_export(df_raw):
        # ── Ximmio bakwagen export ──────────────────────────────
        cats_series = df_raw['SubTaskDesc'].apply(lambda v: parse_subtaskdesc(v)[0])
        cats_upper  = cats_series.fillna('').str.upper()
        overgeslagen_rows = []
        skip_indices = set()
        for idx, row in df_raw.iterrows():
            cat, container = parse_subtaskdesc(str(row.get('SubTaskDesc', '')))
            if cat == 'REMOVE':
                continue
            def leeg(val):
                import pandas as pd
                if val is None:
                    return True
                if isinstance(val, float) and pd.isna(val):
                    return True
                return str(val).strip().lower() in ('', 'nan', 'none')

            redenen = []
            containercode = strip_spaces(container or '')
            streetname    = row.get('Straat', '')
            zipcode_raw   = row.get('Postcode', '')
            city_raw      = row.get('Stad', '')
            huisnummer_v  = row.get('Huisnummer', '')
            subtask_v     = row.get('SubTaskDesc', '')

            streetname_s  = '' if leeg(streetname)  else str(streetname).strip()
            zipcode_s     = '' if leeg(zipcode_raw)  else strip_spaces(str(zipcode_raw))
            city_s        = '' if leeg(city_raw)     else str(city_raw).strip()
            huisnummer_s  = '' if leeg(huisnummer_v) else str(huisnummer_v).strip()
            subtask_s     = '' if leeg(subtask_v)    else str(subtask_v).strip()

            if len(containercode) < 5:
                redenen.append(f"ContainerCode te kort of leeg ('{containercode}')")
            if not streetname_s:
                redenen.append("Straat leeg")
            if not zipcode_s:
                redenen.append("Postcode leeg")
            if not city_s:
                redenen.append("Stad leeg")
            if not huisnummer_s:
                redenen.append("Huisnummer leeg")
            if not subtask_s:
                redenen.append("SubTaskDesc leeg")
            if redenen:
                skip_indices.add(idx)
                huisnummer = str(row.get('Huisnummer', '')).strip()
                overgeslagen_rows.append({
                    'rij':       idx + 2,
                    'adres':     f"{streetname_s} {huisnummer_s}".strip() or '—',
                    'postcode':  zipcode_s or '—',
                    'container': containercode or '—',
                    'reden':     ' · '.join(redenen),
                })

        counts = {
            'wissel':            int((cats_upper == 'CHANGE').sum()),
            'uitzetten':         int(((cats_upper == 'NEW') | (cats_upper == 'EXTRA')).sum()),
            'overgeslagen':      len(overgeslagen_rows),
            'overgeslagen_rows': overgeslagen_rows,
        }

        result_df = dataframe_from_ximmio_export(df_raw, skip_indices=skip_indices)

        if not result_df.empty:
            result_df['_hn_int'] = pd.to_numeric(result_df['_huisnummer_raw'], errors='coerce').fillna(0).astype(int)
            result_df = result_df.sort_values(
                by=['_zipcode_raw', '_hn_int', '_huisletter_raw', '_toevoeging_raw'],
                ascending=True, na_position='last'
            ).reset_index(drop=True)
            result_df = result_df.drop(columns=['_cat', '_zipcode_raw', '_huisnummer_raw',
                                                 '_huisletter_raw', '_toevoeging_raw', '_hn_int'])

        return result_df, counts

    else:
        # ── Standaard formaat (Nederlands of Engels) ───────────
        df = df_raw.copy()

        cols = set(df.columns.str.strip())
        is_nl = 'ContainerCode' in cols and 'Straat' in cols and 'Postcode' in cols

        def safe_col(df, col):
            return df[col].fillna('').astype(str).str.strip() if col in df.columns else pd.Series([''] * len(df))

        def leeg_std(val):
            if val is None: return True
            if isinstance(val, float) and pd.isna(val): return True
            return str(val).strip().lower() in ('', 'nan', 'none')

        if is_nl:
            REQUIRED_NL = ['ContainerCode', 'Straat', 'Postcode', 'Huisnummer', 'Woonplaats']

            overgeslagen_rows = []
            skip_idx = set()
            for idx, row in df.iterrows():
                redenen = []
                containercode = strip_spaces(str(row.get('ContainerCode', '') or ''))
                straat_v   = row.get('Straat', '')
                postcode_v = row.get('Postcode', '')
                hn_v       = row.get('Huisnummer', '')
                woon_v     = row.get('Woonplaats', '')
                straat_s   = '' if leeg_std(straat_v)   else str(straat_v).strip()
                postcode_s = '' if leeg_std(postcode_v) else strip_spaces(str(postcode_v))
                hn_s       = '' if leeg_std(hn_v)       else str(hn_v).strip()
                woon_s     = '' if leeg_std(woon_v)     else str(woon_v).strip()
                if len(containercode) < 5:
                    redenen.append(f"ContainerCode te kort of leeg ('{containercode}')")
                if not straat_s:
                    redenen.append("Straat leeg")
                if not postcode_s:
                    redenen.append("Postcode leeg")
                if not hn_s:
                    redenen.append("Huisnummer leeg")
                if not woon_s:
                    redenen.append("Woonplaats leeg")
                if redenen:
                    skip_idx.add(idx)
                    overgeslagen_rows.append({
                        'rij':       idx + 2,
                        'adres':     f"{straat_s} {hn_s}".strip() or '—',
                        'postcode':  postcode_s or '—',
                        'container': containercode or '—',
                        'reden':     ' · '.join(redenen),
                    })

            counts = {
                'wissel':            0,
                'uitzetten':         len(df) - len(skip_idx),
                'overgeslagen':      len(overgeslagen_rows),
                'overgeslagen_rows': overgeslagen_rows,
            }

            df['_hn_int'] = pd.to_numeric(df.get('Huisnummer'), errors='coerce').fillna(0).astype(int)
            df = df.sort_values(
                by=['Postcode', '_hn_int', 'Huisletter', 'Huisnummertoevoeging'],
                ascending=True, na_position='last'
            ).reset_index(drop=True)

            result_df = pd.DataFrame([
                {
                    'containertype': strip_spaces(str(row.get('ContainerCode', '') or '')),
                    'straat':        str(row.get('Straat', '') or '').strip(),
                    'huisnummer':    str(row.get('Huisnummer', '') or '').strip(),
                    'toevoeging':    (
                        ('' if leeg_std(row.get('Huisletter')) else str(row['Huisletter']).strip()) +
                        ('' if leeg_std(row.get('Huisnummertoevoeging')) else str(row['Huisnummertoevoeging']).strip())
                    ).strip(),
                    'postcode':      strip_spaces(str(row.get('Postcode', '') or '')),
                    'woonplaats':    str(row.get('Woonplaats', '') or '').strip(),
                }
                for idx, row in df.iterrows() if idx not in skip_idx
            ])

        else:
            raise ValueError(
                "Onbekend bestandsformaat. Het bestand moet de Ximmio bakwagen kolommen bevatten "
                "(Stad, Straat, Huisnummer, Postcode, SubTaskDesc) of het Nederlandse standaardformaat "
                "(ContainerCode, Straat, Postcode, Huisnummer, Woonplaats)."
            )

        return result_df, counts


# -------------------------------------------------------
# Routelijst generator
# -------------------------------------------------------

ORDERS_PER_PAGE = 4


def _make_barcode_image(value):
    """Genereer een Code128 barcode PNG op natuurlijke resolutie (geen stretch)."""
    encoder = Code128Encoder(str(value))
    img = Image.open(BytesIO(encoder.get_imagedata()))
    bbox = img.getbbox()
    if bbox:
        img = img.crop((bbox[0], 0, bbox[2], img.height))
    draw = ImageDraw.Draw(img)
    w, h = img.size
    draw.rectangle([0, h - 30, w, h], fill='white')
    img = img.crop((0, 0, w, h - 28))
    buf = BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    w2, h2 = img.size
    natural_w_cm = w2 / 96 * 2.54
    natural_h_cm = h2 / 96 * 2.54
    max_w_cm = 4.8
    if natural_w_cm > max_w_cm:
        scale = max_w_cm / natural_w_cm
        width_cm  = round(max_w_cm, 2)
        height_cm = round(natural_h_cm * scale, 2)
    else:
        width_cm  = round(natural_w_cm, 2)
        height_cm = round(natural_h_cm, 2)
    return buf, width_cm, height_cm


def _set_cell_border(cell, **kwargs):
    """Zet celranden via OxmlElement."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        attrs = kwargs.get(edge, {})
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'),   attrs.get('val', 'single'))
        el.set(qn('w:sz'),    str(attrs.get('sz', 4)))
        el.set(qn('w:color'), attrs.get('color', '000000'))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _make_field_run(paragraph, field_code, font_size_pt=12, bold=False, color=None):
    """Voeg een Word-veld (PAGE, NUMPAGES etc.) inline toe als run.
    color: optionele hex-kleurstring, bijv. '888888' voor grijs.
    """
    def _new_run_with_rpr():
        r = OxmlElement('w:r')
        rpr = OxmlElement('w:rPr')
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(font_size_pt * 2)))
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(font_size_pt * 2)))
        fn = OxmlElement('w:rFonts')
        fn.set(qn('w:ascii'), 'Arial')
        fn.set(qn('w:hAnsi'), 'Arial')
        if bold:
            b = OxmlElement('w:b')
            rpr.append(b)
        if color:
            cl = OxmlElement('w:color')
            cl.set(qn('w:val'), color)
            rpr.append(cl)
        rpr.append(fn)
        rpr.append(sz)
        rpr.append(szCs)
        r.append(rpr)
        return r

    r1 = _new_run_with_rpr()
    fc1 = OxmlElement('w:fldChar')
    fc1.set(qn('w:fldCharType'), 'begin')
    r1.append(fc1)
    paragraph._p.append(r1)

    r2 = _new_run_with_rpr()
    instr = OxmlElement('w:instrText')
    instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instr.text = f' {field_code} '
    r2.append(instr)
    paragraph._p.append(r2)

    r3 = _new_run_with_rpr()
    fc3 = OxmlElement('w:fldChar')
    fc3.set(qn('w:fldCharType'), 'end')
    r3.append(fc3)
    paragraph._p.append(r3)


def _add_right_tab(paragraph, position_dxa):
    """Voeg een rechts-uitgelijnde tabstop toe aan een paragraaf."""
    pPr = paragraph._p.get_or_add_pPr()
    tabs_el = pPr.find(qn('w:tabs'))
    if tabs_el is None:
        tabs_el = OxmlElement('w:tabs')
        pPr.append(tabs_el)
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(position_dxa))
    tabs_el.append(tab)


def _hdr_run(paragraph, text, bold=False, size_pt=12, color=None):
    """Voeg een opgemaakte run toe aan een paragraaf.
    color: tuple (R,G,B) of None.
    """
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = 'Arial'
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def _add_tab_run(paragraph, size_pt=12):
    """Voeg een tab-run toe aan een paragraaf."""
    run = paragraph.add_run()
    run.font.name = 'Arial'
    run.font.size = Pt(size_pt)
    tab_el = OxmlElement('w:tab')
    run._r.append(tab_el)
    return run


def _add_bottom_border_para(paragraph, color='2E75B6', sz=6):
    """Voeg een blauwe lijn toe onder een paragraaf."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    str(sz))
    bottom.set(qn('w:color'), color)
    bottom.set(qn('w:space'), '1')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _fetch_logo_bytes():
    """Laad het Avalex-logo. Probeert eerst lokaal (naast app.py), dan via URL.
    Geeft een verse BytesIO terug of None als het niet lukt.
    """
    import os

    # 1. Probeer lokaal bestand (zelfde map als het script)
    local_paths = [
        os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Avalexlogo.png'),
        'Avalexlogo.png',
    ]
    for path in local_paths:
        if os.path.isfile(path):
            try:
                with open(path, 'rb') as f:
                    data = f.read()
                buf = BytesIO(data)
                Image.open(buf)   # valideer
                buf.seek(0)
                return buf
            except Exception:
                pass

    # 2. Probeer via URL als fallback
    if AVALEX_LOGO_URL and 'your-org' not in AVALEX_LOGO_URL:
        try:
            import urllib.request
            with urllib.request.urlopen(AVALEX_LOGO_URL, timeout=5) as resp:
                data = resp.read()
            buf = BytesIO(data)
            Image.open(buf)   # valideer
            buf.seek(0)
            return buf
        except Exception:
            pass

    return None


def _add_header(section, meta, content_width_dxa):
    """
    Bouw de paginaheader op met twee regels:

    Regel 1 (rechts uitgelijnd, klein grijs):
        Pagina X/Y

    Regel 2 (links → center tab → rechts tab):
        Datum: <waarde>  [CENTER TAB]  Wagen: X   Kenteken: X   Bestuurder: X  [RIGHT TAB]  Route: <waarde>

    + blauwe scheidingslijn onder regel 2.
    """
    header = section.header
    for p in header.paragraphs:
        p.clear()

    center_tab_dxa = content_width_dxa // 2

    # ── Regel 1: Pagina X/Y rechtsboven (klein, grijs) ────────
    p1 = header.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _hdr_run(p1, 'Pagina ', bold=True, size_pt=10, color=(0x88, 0x88, 0x88))
    _make_field_run(p1, 'PAGE',     font_size_pt=10, color='888888')
    _hdr_run(p1, '/', size_pt=10, color=(0x88, 0x88, 0x88))
    _make_field_run(p1, 'NUMPAGES', font_size_pt=10, color='888888')
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after  = Pt(8)

    # ── Regel 2: Datum links | center tab | Wagen/Kenteken/Bestuurder | right tab | Route ──
    p2 = header.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Center tab voor de middelste velden
    pPr = p2._p.get_or_add_pPr()
    tabs_el = OxmlElement('w:tabs')
    tab_c = OxmlElement('w:tab')
    tab_c.set(qn('w:val'), 'center')
    tab_c.set(qn('w:pos'), str(center_tab_dxa))
    tabs_el.append(tab_c)
    tab_r = OxmlElement('w:tab')
    tab_r.set(qn('w:val'), 'right')
    tab_r.set(qn('w:pos'), str(content_width_dxa))
    tabs_el.append(tab_r)
    pPr.append(tabs_el)

    _hdr_run(p2, 'Datum: ', bold=True)
    _hdr_run(p2, meta.get('taakdatum', ''))
    _add_tab_run(p2)   # → center tab
    _hdr_run(p2, 'Wagen: ', bold=True)
    _hdr_run(p2, meta.get('wagen', '') + '   ')
    _hdr_run(p2, 'Kenteken: ', bold=True)
    _hdr_run(p2, meta.get('voertuig', '') + '   ')
    _hdr_run(p2, 'Bestuurder: ', bold=True)
    _hdr_run(p2, meta.get('naambestuurder', ''))
    _add_tab_run(p2)   # → right tab
    _hdr_run(p2, 'Route: ', bold=True)
    _hdr_run(p2, meta.get('route', ''))

    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(2)

    # Blauwe lijn onder regel 2
    _add_bottom_border_para(p2)


def _add_footer(section, logo_buf):
    """
    Voeg een footer toe met het Avalex-logo gecentreerd onderaan de pagina.
    logo_buf: BytesIO met PNG-data, of None (dan wordt de footer weggelaten).
    """
    if logo_buf is None:
        return

    footer = section.footer
    for p in footer.paragraphs:
        p.clear()

    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)

    run = p.add_run()
    logo_buf.seek(0)
    # Logo hoogte: 1.2 cm, breedte proportioneel
    run.add_picture(logo_buf, height=Cm(1.2))


def generate_routelijst(df_raw, meta):
    """
    Genereer een landschaps-A4 routelijst Word-document.

    meta = {
        'taakdatum':      str,
        'wagen':          str,
        'voertuig':       str,
        'naambestuurder': str,
        'route':          str,
    }
    """
    doc = Document()

    # ── Pagina-instellingen: landscape A4 ──────────────────────
    section = doc.sections[0]
    # A4 landscape: stel in als portrait A4 + orient=landscape
    # zodat Word het herkent als standaard A4 en niet als 'Aangepast formaat'
    section.page_width    = Cm(21.0)
    section.page_height   = Cm(29.7)
    # Landscape-oriëntatie instellen via XML (python-docx heeft geen native API hiervoor)
    from docx.oxml import OxmlElement as _OxmlElement
    from docx.oxml.ns import qn as _qn
    sectPr = section._sectPr
    pgSz = sectPr.find(_qn('w:pgSz'))
    if pgSz is None:
        pgSz = _OxmlElement('w:pgSz')
        sectPr.insert(0, pgSz)
    pgSz.set(_qn('w:orient'), 'landscape')
    # Zet breedte/hoogte ook direct in DXA (1 cm = 567 DXA) voor zekerheid
    pgSz.set(_qn('w:w'), str(int(29.7 * 567)))   # 16834
    pgSz.set(_qn('w:h'), str(int(21.0 * 567)))   # 11907
    section.left_margin   = Cm(1.5)
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(3.8)
    section.bottom_margin = Cm(1.5)

    # Contentbreedte in DXA: (29.7 - 1.5 - 1.5) * 567 ≈ 15138
    CONTENT_WIDTH_DXA = int((29.7 - 1.5 - 1.5) * 567)

    # Logo ophalen (één keer, hergebruiken voor alle pagina's)
    logo_buf = _fetch_logo_bytes()

    # ── Header en footer instellen ──────────────────────────────
    _add_header(section, meta, CONTENT_WIDTH_DXA)
    _add_footer(section, logo_buf)

    # ── Orders verwerken ────────────────────────────────────────
    rows_data = []
    for _, row in df_raw.iterrows():
        subtask_raw = str(row.get('SubTaskDesc', '') or '')
        cat, container = parse_subtaskdesc(subtask_raw)

        straat   = str(row.get('Straat', '') or '').strip()
        hn       = str(row.get('Huisnummer', '') or '').strip()
        hl_raw   = row.get('Huisletter', '')
        tv_raw   = row.get('Huisnummer toevoeging', '')
        hl       = '' if pd.isna(hl_raw) else str(hl_raw).strip()
        tv       = '' if pd.isna(tv_raw) else str(tv_raw).strip()
        postcode = strip_spaces(str(row.get('Postcode', '') or ''))
        stad     = str(row.get('Stad', '') or '').strip()

        adres_r1 = f"{straat} {hn} {hl}{tv}".strip()
        adres_r2 = f"{postcode}  {stad}"

        subtaak_id  = row.get('SubtaakID', '')
        barcode_val = str(int(subtaak_id)) if pd.notna(subtaak_id) and str(subtaak_id).strip() not in ('', 'nan') else subtask_raw[:20]

        rows_data.append({
            'barcode_val':  barcode_val,
            'adres_r1':     adres_r1,
            'adres_r2':     adres_r2,
            'subtaskdesc':  subtask_raw,
            'toelichting':  '' if pd.isna(row.get('Toelichting', '')) else str(row.get('Toelichting', '') or '').strip(),
            'is_remove':    cat == 'REMOVE',
            '_postcode':    postcode,
            '_hn_int':      int(hn) if hn.isdigit() else 0,
            '_hl':          hl,
            '_tv':          tv,
        })

    rows_data.sort(key=lambda r: (r['_postcode'], r['_hn_int'], r['_hl'], r['_tv']))

    COL_W = [Cm(5.5), Cm(11.0), Cm(10.2)]

    def _blank_cell_run(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = 'Arial'
        return p

    def _add_para(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
        p = cell.add_paragraph()
        p.alignment = align
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = 'Arial'
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        return p

    pages = [rows_data[i:i+ORDERS_PER_PAGE] for i in range(0, max(len(rows_data), 1), ORDERS_PER_PAGE)]

    for page_idx, page_orders in enumerate(pages):
        if page_idx > 0:
            doc.add_page_break()

        tbl = doc.add_table(rows=1 + len(page_orders), cols=3)
        tbl.style = 'Table Grid'
        tbl.autofit = False

        # ── Koptekstrij ────────────────────────────────────────
        hdr_row = tbl.rows[0]
        for ci, cell in enumerate(hdr_row.cells):
            cell.width = COL_W[ci]
            p = cell.paragraphs[0]
            run = p.add_run(['Adres / Barcode', 'Omschrijving / Opmerkingen', 'Containersticker'][ci])
            run.bold = True
            run.font.size = Pt(8)
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '2E75B6')
            tcPr.append(shd)

        for row_i, order in enumerate(page_orders):
            row = tbl.rows[1 + row_i]
            for ci, cell in enumerate(row.cells):
                cell.width = COL_W[ci]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER if ci in (0, 2) else WD_ALIGN_VERTICAL.TOP
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after  = Pt(0)

            # ── Cel 0: barcode + adres ──────────────────────────
            c0 = row.cells[0]
            try:
                bc_buf, bc_w, bc_h = _make_barcode_image(order['barcode_val'])
                p_bc = c0.paragraphs[0]
                p_bc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_bc = p_bc.add_run()
                run_bc.add_picture(bc_buf, width=Cm(bc_w), height=Cm(bc_h))
                p_bc.paragraph_format.space_before = Pt(2)
                p_bc.paragraph_format.space_after  = Pt(0)
            except Exception:
                _blank_cell_run(c0, order['barcode_val'], bold=True)

            p_id = c0.add_paragraph()
            p_id.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_id.paragraph_format.space_before = Pt(0)
            p_id.paragraph_format.space_after  = Pt(6)
            r_id = p_id.add_run(order['barcode_val'])
            r_id.font.size = Pt(7)
            r_id.font.name = 'Arial'
            r_id.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

            _add_para(c0, order['adres_r1'], bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
            _add_para(c0, order['adres_r2'], size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

            # ── Cel 1: SubTaskDesc + opmerkingen ───────────────
            c1 = row.cells[1]
            _blank_cell_run(c1, 'Omschrijving:', bold=True, size=8)
            import textwrap as _tw
            # Wrap lange tekst over meerdere losse paragrafen (\n werkt niet in python-docx)
            for desc_line in _tw.wrap(order['subtaskdesc'], width=65) or ['']:
                _add_para(c1, desc_line, size=9)
            if order['toelichting']:
                _add_para(c1, f"Toelichting: {order['toelichting']}", size=8)
            _add_para(c1, 'Opmerkingen:', bold=True, size=8)
            for _ in range(3):
                lijn_p = c1.add_paragraph()
                lijn_p.paragraph_format.space_before = Pt(0)
                lijn_p.paragraph_format.space_after  = Pt(2)
                pPr = lijn_p._p.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                bot = OxmlElement('w:bottom')
                bot.set(qn('w:val'),   'single')
                bot.set(qn('w:sz'),    '4')
                bot.set(qn('w:color'), 'AAAAAA')
                bot.set(qn('w:space'), '2')
                pBdr.append(bot)
                pPr.append(pBdr)
                lijn_p.add_run('').font.size = Pt(10)

            # ── Cel 2: containersticker / innemen ───────────────
            c2 = row.cells[2]
            p_sticker = c2.paragraphs[0]
            p_sticker.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_sticker.paragraph_format.space_before = Pt(8)
            if order['is_remove']:
                run_st = p_sticker.add_run('📥 INNEMEN')
                run_st.font.size = Pt(14)
                run_st.bold = True
                run_st.font.name = 'Arial'
                run_st.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
            else:
                run_st = p_sticker.add_run('<CONTAINERSTICKER>')
                run_st.font.size = Pt(9)
                run_st.font.name = 'Arial'
                run_st.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer


# -------------------------------------------------------
# Streamlit UI
# -------------------------------------------------------

st.set_page_config(page_title="Labelgenerator", page_icon="📦")
st.title("📦 Containerlabelgenerator")

tab_xlsx, tab_manual = st.tabs(["📂 XLSX uploaden", "✏️ Handmatig invoeren"])

# ── Tab 1: XLSX upload ─────────────────────────────────
with tab_xlsx:
    st.write("Upload een XLSX bestand om labels te genereren met barcodes.")

    uploaded_file = st.file_uploader("Sleep je .xlsx bestand hiernaartoe", type=["xlsx"])

    if uploaded_file:
        if st.button("Verwerken", key="btn_xlsx"):
            with st.spinner("Bezig met verwerken..."):
                try:
                    uploaded_file.seek(0)
                    df_raw_upload = pd.read_excel(uploaded_file)
                    is_ximmio = is_ximmio_export(df_raw_upload)
                    uploaded_file.seek(0)

                    df, counts = dataframe_from_file(uploaded_file)
                    docx_labels = generate_word_from_dataframe(df)
                    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M")

                    st.success(f"✅ {len(df)} label(s) gegenereerd!")

                    col1, col2, col3 = st.columns(3)
                    col1.metric("🔄 Wissel",       counts['wissel'],      help="CategoryName = CHANGE")
                    col2.metric("📦 Uitzetten",    counts['uitzetten'],   help="CategoryName = NEW of EXTRA")
                    col3.metric("⛔ Overgeslagen", counts['overgeslagen'], help="Ontbrekende/ongeldige velden")

                    if counts['overgeslagen_rows']:
                        with st.expander(f"⛔ {counts['overgeslagen']} overgeslagen rij(en) — klik om te bekijken"):
                            overgeslagen_df = pd.DataFrame(counts['overgeslagen_rows'])
                            overgeslagen_df.columns = ['Rij', 'Adres', 'Postcode', 'Containertype', 'Reden']
                            st.dataframe(overgeslagen_df, hide_index=True, width="stretch")

                    st.markdown("---")

                    if is_ximmio:
                        first_row = df_raw_upload.iloc[0]
                        def _s(col):
                            v = first_row.get(col, '')
                            return '' if pd.isna(v) else str(v).strip()
                        meta = {
                            'taakdatum':      _s('Taakdatum'),
                            'wagen':          _s('Wagen'),
                            'voertuig':       _s('Kenteken') or _s('Voertuig'),
                            'naambestuurder': _s('NaamBestuurder'),
                            'route':          _s('TaskDesc'),
                        }
                        docx_route = generate_routelijst(df_raw_upload, meta)
                        n_orders = sum(
                            1 for _, row in df_raw_upload.iterrows()
                            if parse_subtaskdesc(str(row.get('SubTaskDesc', '')))[0] is not None
                        )
                        n_pages = -(-n_orders // ORDERS_PER_PAGE)

                        taakdatum_raw = meta.get('taakdatum', '') or timestamp
                        route_raw     = meta.get('route', '') or 'route'
                        datum_clean = re.sub(r'[^0-9]', '', str(taakdatum_raw))[:8] or timestamp[:8]
                        route_clean = re.sub(r'[^a-zA-Z0-9_]', '_', route_raw).strip('_')
                        bestand_prefix = f"{datum_clean}_{route_clean}"

                        col_dl1, col_dl2 = st.columns(2)
                        with col_dl1:
                            st.download_button(
                                label="📥 Download containerlabels",
                                data=docx_labels,
                                file_name=f"{bestand_prefix}_containerlabels.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key="dl_labels",
                                use_container_width=True,
                            )
                        with col_dl2:
                            st.download_button(
                                label=f"📋 Download routelijst ({n_orders} orders, {n_pages} pag.)",
                                data=docx_route,
                                file_name=f"{bestand_prefix}_routelijst.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key="dl_route",
                                use_container_width=True,
                            )
                    else:
                        st.download_button(
                            label="📥 Download containerlabels",
                            data=docx_labels,
                            file_name=f"containerlabels_{timestamp}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="dl_xlsx"
                        )
                except ValueError as e:
                    st.error(str(e))
                except Exception as e:
                    import traceback
                    st.error(f"Fout bij verwerken: {e}")
                    st.code(traceback.format_exc(), language="text")

# ── Tab 2: Handmatig invoeren ──────────────────────────
with tab_manual:
    st.write(f"Vul hieronder handmatig de gegevens in (maximaal {MAX_ROWS} labels).")

    if 'num_rows' not in st.session_state:
        st.session_state.num_rows = 1

    st.session_state.num_rows = max(1, min(st.session_state.num_rows, MAX_ROWS))

    col_add, col_remove = st.columns([1, 1])
    with col_add:
        if st.button("➕ Rij toevoegen", disabled=st.session_state.num_rows >= MAX_ROWS):
            st.session_state.num_rows = min(st.session_state.num_rows + 1, MAX_ROWS)
    with col_remove:
        if st.button("➖ Rij verwijderen", disabled=st.session_state.num_rows <= 1):
            st.session_state.num_rows = max(st.session_state.num_rows - 1, 1)

    st.markdown("---")

    rows = []
    for i in range(st.session_state.num_rows):
        st.markdown(f"**Label {i + 1}**")
        c1, c2, c3, c4, c5, c6 = st.columns([2, 2, 1, 1, 1.5, 2])
        with c1:
            containertype = st.text_input("Containertype", key=f"ct_{i}", placeholder="OPK_140L")
        with c2:
            straat = st.text_input("Straat", key=f"st_{i}", placeholder="Teststraat")
        with c3:
            huisnummer = st.text_input("Nr.", key=f"hn_{i}", placeholder="9")
        with c4:
            toevoeging = st.text_input("Toev.", key=f"tv_{i}", placeholder="A")
        with c5:
            postcode = st.text_input("Postcode", key=f"pc_{i}", placeholder="1234AA")
        with c6:
            woonplaats = st.text_input("Woonplaats", key=f"wp_{i}", placeholder="Rijswijk")

        rows.append({
            'containertype': strip_spaces(containertype),
            'straat':        straat.strip(),
            'huisnummer':    huisnummer.strip(),
            'toevoeging':    toevoeging.strip(),
            'postcode':      strip_spaces(postcode),
            'woonplaats':    woonplaats.strip(),
        })

    st.markdown("---")

    if st.button("Verwerken", key="btn_manual"):
        df_manual = pd.DataFrame(rows)
        df_manual = df_manual[
            (df_manual['postcode'] != '') & (df_manual['huisnummer'] != '')
        ].reset_index(drop=True)

        if df_manual.empty:
            st.warning("Vul minimaal postcode en huisnummer in voor één label.")
        else:
            with st.spinner("Bezig met verwerken..."):
                docx_file = generate_word_from_dataframe(df_manual)
                timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M")
                st.success(f"{len(df_manual)} label(s) gegenereerd!")
                st.download_button(
                    label="📥 Download Word-bestand",
                    data=docx_file,
                    file_name=f"containerlabels_{timestamp}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_manual"
                )